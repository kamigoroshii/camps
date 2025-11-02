"""
Document processing utilities for extracting text from various file formats
"""

import logging
from typing import List
from io import BytesIO

import PyPDF2
import docx
import pandas as pd
from PIL import Image
import pytesseract

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document processing and text extraction"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n[Page {page_num}]\n{page_text}\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_xlsx(file_content: bytes) -> str:
        """Extract text from Excel file"""
        try:
            df = pd.read_excel(BytesIO(file_content), sheet_name=None)
            text = ""
            
            for sheet_name, sheet_df in df.items():
                text += f"\n[Sheet: {sheet_name}]\n"
                text += sheet_df.to_string(index=False) + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting Excel: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_csv(file_content: bytes) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(BytesIO(file_content))
            text = f"Table Data:\n{df.to_string(index=False)}"
            return text
        except Exception as e:
            logger.error(f"Error extracting CSV: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_image(file_content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            if not settings.ENABLE_OCR:
                return ""
            
            image = Image.open(BytesIO(file_content))
            text = pytesseract.image_to_string(image, lang=settings.TESSERACT_LANG)
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            return ""
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into overlapping chunks"""
        if not text.strip():
            return []
        
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP
        
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = " ".join(chunk_words)
            if chunk_text.strip():
                chunks.append(chunk_text)
            
            # Limit number of chunks per file
            if len(chunks) >= settings.MAX_CHUNKS_PER_FILE:
                logger.warning(f"Reached maximum chunks limit: {settings.MAX_CHUNKS_PER_FILE}")
                break
        
        return chunks
    
    @staticmethod
    def extract_text(file_content: bytes, file_extension: str) -> str:
        """Extract text based on file type"""
        ext = file_extension.lower()
        
        if ext == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_content)
        elif ext == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_content)
        elif ext == '.txt':
            return file_content.decode('utf-8', errors='ignore')
        elif ext == '.xlsx':
            return DocumentProcessor.extract_text_from_xlsx(file_content)
        elif ext == '.csv':
            return DocumentProcessor.extract_text_from_csv(file_content)
        elif ext in ['.jpg', '.jpeg', '.png']:
            return DocumentProcessor.extract_text_from_image(file_content)
        else:
            logger.error(f"Unsupported file type: {ext}")
            return ""
